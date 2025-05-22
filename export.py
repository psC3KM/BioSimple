import csv
import json
from typing import List
from openpyxl import Workbook
from openpyxl.styles import Font
from openpyxl.utils import get_column_letter
from docx import Document
from sample import BiologicalSample


def export_to_excel (samples, filename='samples.xlsx'):
    """
    Экспорт списка образцов в файл Excel.
    Создание списка объектов BiologicalSample.
    Назначение имя файла для сохранения.
    """
    wb = Workbook()
    ws = wb.active
    ws.title = 'Биологические образцы'
    
    headers = ['ID','Название','Тип','Дата сбора','Место','Примечания']
    ws.append(headers)
    
    for cell in ws[1]:
        cell.font = Font(bold=True)
    
    for sample in samples:
        row = [
            sample.id,
            sample.name,
            sample.type,
            sample.collection_date,
            sample.location,
            sample.notes
        ]
        ws.append(row)
        
    for col in ws.columns:
        max_length = 0
        column = col[0].column_letter
        for cell in col:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(str(cell.value))
            except:
                pass
        adjusted_width = (max_length + 2)
        ws.column_dimensions[column].width = adjusted_width

    wb.save(filename)


def export_to_word(samples:List[BiologicalSample],filename:str="samples.docx"):
    """
    Экспорт списка образцов в файл Docx.
    Создаёт документ с карточками образцов, где для каждого образца указаны:
    - ID и название (как заголовок)
    - Тип
    - Дата отбора
    - Место отбора
    - Примечение (если имеется)
    """
    doc = Document()
    doc.add_heading('Карточки биологических образцов', level=1)
    for sample in samples:
        doc.add_heading(f'Образец #{sample.id}: {sample.name}', level=2)
        doc.add_paragraph(f'Тип: {sample.type}')
        doc.add_paragraph(f'Дата отбора: {sample.collection_date}')
        doc.add_paragraph(f'Место: {sample.location}')
        
        if sample.notes:
            doc.add_paragraph(f'Примечание: {sample.notes}')
            
        doc.add_paragraph('-' * 40)
    
    doc.save(filename)


def export_to_csv(samples:List[BiologicalSample],filename:str='samples.csv'):
    with open(filename,'w',newline='',encoding='utf-8') as file:
        writer = csv.writer(file)
        # Запишите заголовки и данные
        for writer in writers:
            row = [
                sample.id,
                sample.name,
                sample.type,
                sample.collection_date,
                sample.location,
                sample.notes
            ]
        writer.append()
    
    writer.save()
